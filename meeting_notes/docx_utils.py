from typing import List


def extract_text_from_docx(path: str) -> str:
    """Extract readable text from a Teams transcript .docx.

    - Collects paragraphs.
    - Flattens tables, mapping typical columns (Timestamp, Speaker, Transcript)
      into lines like: "[12:34] Alice: content".
    - Skips empty lines and common header rows.
    """
    try:
        import docx  # type: ignore
    except ImportError as ie:
        raise RuntimeError(
            "python-docx is not installed. Run 'pip install python-docx' or 'pip install -r requirements.txt'."
        ) from ie

    try:
        document = docx.Document(path)
    except Exception as e:
        raise RuntimeError(f"Failed to read DOCX file: {e}") from e

    lines: List[str] = []

    # Extract paragraphs
    for p in document.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    # Extract tables (common for Teams transcripts)
    header_tokens = {"speaker", "timestamp", "time", "transcript", "utterance"}
    for tbl in document.tables:
        for ridx, row in enumerate(tbl.rows):
            cells = [((c.text or "").strip()) for c in row.cells]
            if not any(cells):
                continue
            # Skip header row heuristically
            if ridx == 0 and any(cell.lower() in header_tokens for cell in cells):
                continue
            if len(cells) >= 3:
                # Assume columns: Timestamp | Speaker | Transcript
                timestamp, speaker, utterance = cells[0], cells[1], cells[2]
                prefix = f"[{timestamp}] {speaker}: " if timestamp else f"{speaker}: "
                lines.append(prefix + utterance)
            elif len(cells) == 2:
                # Speaker | Transcript (no timestamp)
                lines.append(f"{cells[0]}: {cells[1]}")
            else:
                lines.append(cells[0])

    # Normalize by removing empty lines
    cleaned = [ln for ln in lines if ln.strip()]
    return "\n".join(cleaned)